from __future__ import annotations

import html
import io
import json
import os
import re
import shutil
import subprocess
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any

from qa_pipeline.agents.existing_framework_control.deep_framework_agents import build_deep_framework_understanding
from qa_pipeline.agents.existing_framework_control.structure_discovery import build_structure_profile
from qa_pipeline.core.active_context import write_active_context
from qa_pipeline.core.io import read_json, write_json
from qa_pipeline.core.paths import QA_CACHE_DIR, REPO_ROOT, feature_testcase_path
from qa_pipeline.core.runtime_logger import log_event
from qa_pipeline.core.text import camel_case, pascal_case, safe_id
from qa_pipeline.agents.phase2_source_intake_rag.ingest import write_functional_testcases_markdown

_TEXT_EXTENSIONS = {'.txt', '.md', '.markdown', '.json', '.csv', '.yaml', '.yml', '.feature'}
_CODE_EXTENSIONS = {'.ts', '.tsx', '.js', '.jsx', '.mjs', '.cjs'}


def _safe_feature(value: str) -> str:
    return re.sub(r'[^a-z0-9_-]+', '_', (value or 'feature').strip().lower()).strip('_') or 'feature'


def _rel(path: Path, root: Path) -> str:
    return path.resolve().relative_to(root.resolve()).as_posix()


def _resolve_under_root(root: Path, value: str, *, allow_missing: bool = True) -> Path | None:
    raw = (value or '').strip().strip('"').strip("'")
    if not raw:
        return None
    candidate = Path(raw).expanduser()
    if not candidate.is_absolute():
        candidate = root / candidate
    candidate = candidate.resolve()
    try:
        candidate.relative_to(root.resolve())
    except ValueError as exc:
        raise ValueError(f'Path must remain inside the selected framework: {raw}') from exc
    if not allow_missing and not candidate.exists():
        raise ValueError(f'Path does not exist inside the selected framework: {raw}')
    return candidate


def _infer_action(text: str, keyword: str = '') -> str:
    low = f'{keyword} {text}'.lower()
    if keyword.lower() == 'then' or re.search(r'\b(?:verify|validate|expect|expects|should|assert|confirm that|visible|displayed|shown)\b', low):
        return 'verify'
    if re.search(r'https?://|\b(?:base url|application url|test environment url|new browser session)\b', low):
        return 'goto'
    if re.search(r'\b(?:enter|enters|fill|fills|type|types|input|inputs|provide|provides)\b', low):
        return 'fill'
    if re.search(r'\b(?:select|selects|choose|chooses|pick|picks)\b', low) and re.search(r'\b(?:dropdown|combobox|list|status|date|time|type|product|source|interest|option as| as )\b', low):
        return 'select'
    if re.search(r'\b(?:click|clicks|tap|taps|press|presses|submit|submits|confirm|confirms|open|opens|navigate|navigates|go to|launch|launches|visit|visits|select|selects|choose|chooses)\b', low):
        return 'click'
    return 'perform'


def _clean_step_text(value: str) -> str:
    value = re.sub(r'^\s*(?:step\s*)?\d+[\).:\-\s]+', '', value or '', flags=re.I)
    value = re.sub(r'[*_`]+', '', value)
    return value.strip(' \t-•')



def _semantic_target(value: str) -> str:
    text = _clean_step_text(value or 'element')
    fill_in = re.match(r'^(?:in|on)\s+(.+?),\s*(?:enter|type|fill|input|provide)\b', text, re.I)
    if fill_in:
        text = fill_in.group(1)
    text = re.sub(r'^(?:click|tap|press|enter|fill|type|input|provide|select|choose|verify|validate|assert|confirm|open|navigate to|go to)\s+', '', text, flags=re.I)
    text = re.sub(r'^(?:on|the)\s+', '', text, flags=re.I)
    text = re.sub(r'\s+(?:as|with)\s+["\'`]?[^"\'`]+["\'`]?\s*$', '', text, flags=re.I)
    text = re.sub(r'\s+from\s+(?:the\s+)?dropdown\s*$', '', text, flags=re.I)
    text = re.sub(r'\s+(?:and\s+)?(?:click|press|select|open)\s+.+$', '', text, flags=re.I)
    text = re.sub(r'\s+(?:is|should be|must be)\s+(?:displayed|visible|shown|enabled|available|present)\.?$', '', text, flags=re.I)
    return text.strip() or 'element'


def _infer_page_context(title: str, steps: list[dict[str, Any]], fallback: str = 'Home') -> str:
    corpus = ' '.join([title] + [str(step.get('target') or '') for step in steps]).lower()
    if 'salesforce' in corpus or 'customer central' in corpus or 'work queue' in corpus:
        return 'Salesforce'
    if 'loyalty management' in corpus:
        return 'LoyaltyManagement'
    if re.search(r'\blogin|sign in|username|password\b', corpus):
        return 'Login'
    return fallback or 'Home'


def _environment_name_for_step(step_text: str) -> str:
    low = step_text.lower()
    if 'password' in low:
        return 'SALESFORCE_PASSWORD' if 'salesforce' in low or 'store co-worker' in low else 'TEST_PASSWORD'
    if any(token in low for token in ('username', 'user name', 'email')):
        return 'SALESFORCE_USERNAME' if 'salesforce' in low or 'store co-worker' in low or 'agent' in low or 'admin user' in low else 'TEST_USERNAME'
    if 'verification code' in low or re.search(r'\botp\b', low):
        return 'SALESFORCE_VERIFICATION_CODE'
    semantic = re.sub(r'[^A-Za-z0-9]+', '_', _semantic_target(step_text)).strip('_').upper()
    return (semantic[:64] or 'TEST_DATA')


def _extract_test_data(step_text: str, raw_value: str) -> dict[str, Any]:
    raw = (raw_value or '').strip()
    if not raw or raw.lower() in {'n/a', 'na', 'none', 'null', '-'}:
        return {}
    low_step = step_text.lower()
    if any(token in low_step for token in ('password', 'username', 'user name', 'verification code', ' api token', 'token')):
        return {
            'value_env': _environment_name_for_step(step_text),
            'value_sensitive': True,
            'value_source': 'environment_only',
        }
    url_match = re.search(r'https?://[^\s]+', raw)
    if url_match:
        return {'value': url_match.group(0).rstrip('.,;'), 'value_source': 'spreadsheet_test_data'}
    candidate = raw
    labelled = re.match(r'^\s*[^:=]{1,80}\s*[:=]\s*(.+?)\s*$', raw)
    if labelled:
        candidate = labelled.group(1).strip()
    if re.search(r'\b(?:asks?|prompted|displayed|page|message|appears?|successfully)\b', candidate, re.I) and not re.search(r'[=@]|\d{3,}', candidate):
        return {}
    generic = candidate.strip().lower()
    if generic in {'date', 'time', 'product', 'appointment type', 'notes', 'status', 'type'}:
        return {
            'value_env': _environment_name_for_step(step_text),
            'value_source': 'environment_placeholder',
        }
    return {'value': candidate[:500], 'value_source': 'spreadsheet_test_data'}

def _scenario_id(feature: str, index: int, supplied: str = '') -> str:
    if supplied.strip():
        return re.sub(r'[^A-Za-z0-9_-]+', '-', supplied.strip()).strip('-').upper()
    return f'{feature.upper()}-TC-{index:03d}'


def _step_from_text(text: str, page: str = 'Home', keyword: str = '') -> dict[str, Any]:
    clean = _clean_step_text(text)
    expected = ''
    if '=>' in clean:
        clean, expected = [x.strip() for x in clean.split('=>', 1)]
    return {
        'action': _infer_action(clean, keyword),
        'target': clean[:500],
        'page': page or 'Home',
        **({'expected': expected[:500]} if expected else {}),
        **({'gherkin_keyword': keyword.title()} if keyword else {}),
    }


def _parse_json_payload(raw: str, feature: str) -> dict[str, Any] | None:
    try:
        data = json.loads(raw)
    except Exception:
        return None
    if isinstance(data, dict) and isinstance(data.get('scenarios'), list):
        payload = dict(data)
        payload.setdefault('feature', feature)
        payload.setdefault('source_type', 'module2_uploaded')
        return payload
    if isinstance(data, list):
        return {'feature': feature, 'source_type': 'module2_uploaded', 'scenarios': data}
    return None


def _parse_examples(lines: list[str], start: int) -> tuple[list[dict[str, str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        rows.append([c.strip() for c in line.strip('|').split('|')])
        i += 1
    if len(rows) < 2:
        return [], i
    headers = rows[0]
    return [dict(zip(headers, row)) for row in rows[1:] if len(row) == len(headers)], i


def _substitute_outline(value: Any, example: dict[str, str]) -> Any:
    if isinstance(value, str):
        for key, replacement in example.items():
            value = value.replace(f'<{key}>', replacement)
        return value
    if isinstance(value, list):
        return [_substitute_outline(v, example) for v in value]
    if isinstance(value, dict):
        return {k: _substitute_outline(v, example) for k, v in value.items()}
    return value


def parse_gherkin(raw: str, feature: str) -> dict[str, Any]:
    lines = [x.rstrip() for x in (raw or '').splitlines()]
    feature_title = feature.replace('_', ' ').title()
    background: list[dict[str, Any]] = []
    scenarios: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    tags: list[str] = []
    section = ''
    last_primary_keyword = ''
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped or stripped.startswith('#'):
            i += 1
            continue
        if stripped.startswith('@'):
            tags = stripped.split()
            i += 1
            continue
        m = re.match(r'^Feature\s*:\s*(.+)$', stripped, re.I)
        if m:
            feature_title = m.group(1).strip()
            i += 1
            continue
        if re.match(r'^Background\s*:', stripped, re.I):
            section = 'background'
            current = None
            i += 1
            continue
        m = re.match(r'^Scenario(?: Outline| Template)?\s*:\s*(.+)$', stripped, re.I)
        if m:
            current = {
                'id': '',
                'title': m.group(1).strip(),
                'feature': feature,
                'page': pascal_case(feature_title) or 'Home',
                'priority': 'medium',
                'preconditions': [],
                'steps': [],
                'expected_result': '',
                'tags': list(tags),
                'gherkin_source': True,
                'outline': bool(re.match(r'^Scenario(?: Outline| Template)', stripped, re.I)),
            }
            tags = []
            scenarios.append(current)
            section = 'scenario'
            i += 1
            continue
        if re.match(r'^Examples\s*:', stripped, re.I) and current is not None:
            examples, next_i = _parse_examples(lines, i + 1)
            current['examples'] = examples
            i = next_i
            continue
        step_match = re.match(r'^(Given|When|Then|And|But|\*)\s+(.+)$', stripped, re.I)
        if step_match:
            keyword, text = step_match.group(1).title(), step_match.group(2).strip()
            effective_keyword = last_primary_keyword if keyword in {'And', 'But', '*'} and last_primary_keyword else keyword
            if keyword in {'Given', 'When', 'Then'}:
                last_primary_keyword = keyword
            step = _step_from_text(text, current.get('page', 'Home') if current else 'Home', effective_keyword)
            step['gherkin_keyword'] = keyword
            if section == 'background' or current is None:
                background.append(step)
            else:
                current['steps'].append(step)
                if keyword == 'Then':
                    current['expected_result'] = text
            i += 1
            continue
        i += 1

    expanded: list[dict[str, Any]] = []
    for scenario in scenarios:
        steps = [dict(x) for x in background] + list(scenario.get('steps') or [])
        scenario['steps'] = steps
        examples = scenario.pop('examples', [])
        outline = scenario.pop('outline', False)
        if outline and examples:
            for row_no, example in enumerate(examples, 1):
                clone = _substitute_outline(scenario, example)
                clone['title'] = f"{clone.get('title')} [{row_no}]"
                clone['example_data'] = example
                expanded.append(clone)
        else:
            expanded.append(scenario)
    for idx, scenario in enumerate(expanded, 1):
        scenario['id'] = _scenario_id(feature, idx, scenario.get('id', ''))
        if not scenario.get('expected_result'):
            then_step = next((s for s in reversed(scenario.get('steps') or []) if s.get('action') == 'verify'), None)
            scenario['expected_result'] = (then_step or {}).get('target') or 'Expected behavior described by the BDD scenario should be verified.'
    return {
        'feature': feature,
        'feature_title': feature_title,
        'source_type': 'gherkin_bdd',
        'source_format': 'feature',
        'scenario_count': len(expanded),
        'scenarios': expanded,
    }


def _split_plain_blocks(raw: str) -> list[list[str]]:
    lines = [x.rstrip() for x in (raw or '').splitlines()]
    blocks: list[list[str]] = []
    current: list[str] = []
    heading = re.compile(r'^\s*(?:#{1,6}\s*)?(?:test\s*case|testcase|tc|scenario)\s*(?:id|no|number|#)?\s*[:\-]?\s*[A-Za-z0-9_-]+', re.I)
    divider = re.compile(r'^\s*(?:-{3,}|={3,})\s*$')
    for line in lines:
        if heading.match(line) and current:
            blocks.append(current)
            current = [line]
        elif divider.match(line) and current:
            blocks.append(current)
            current = []
        else:
            current.append(line)
    if current:
        blocks.append(current)
    nonempty = [b for b in blocks if any(x.strip() for x in b)]
    return nonempty or [lines]


def _label_value(line: str, labels: tuple[str, ...]) -> str:
    joined = '|'.join(re.escape(x) for x in labels)
    m = re.match(rf'^\s*(?:{joined})\s*[:\-]\s*(.*)$', line, re.I)
    return m.group(1).strip() if m else ''


def _scenario_from_plain_block(block: list[str], feature: str, index: int) -> dict[str, Any]:
    scenario_id = ''
    title = ''
    page = 'Home'
    priority = 'medium'
    preconditions: list[str] = []
    steps: list[dict[str, Any]] = []
    expected_parts: list[str] = []
    section = ''
    for raw_line in block:
        line = raw_line.strip()
        if not line:
            continue
        value = _label_value(line, ('test case id', 'testcase id', 'tc id', 'id'))
        if value:
            scenario_id = value
            continue
        value = _label_value(line, ('test case', 'testcase', 'title', 'scenario', 'test name'))
        if value:
            title = value
            continue
        value = _label_value(line, ('module', 'page', 'screen'))
        if value:
            page = pascal_case(value) or 'Home'
            continue
        value = _label_value(line, ('priority',))
        if value:
            priority = value.lower()
            continue
        if re.match(r'^\s*(preconditions?|prerequisites?)\s*:?\s*$', line, re.I):
            section = 'preconditions'
            continue
        value = _label_value(line, ('precondition', 'preconditions', 'prerequisite'))
        if value:
            preconditions.append(value)
            section = 'preconditions'
            continue
        if re.match(r'^\s*(steps?|test steps?|actions?)\s*:?\s*$', line, re.I):
            section = 'steps'
            continue
        if re.match(r'^\s*(expected results?|expected outcome|result)\s*:?\s*$', line, re.I):
            section = 'expected'
            continue
        value = _label_value(line, ('expected result', 'expected outcome', 'expected'))
        if value:
            expected_parts.append(value)
            section = 'expected'
            continue
        if re.match(r'^(Given|When|Then|And|But)\s+', line, re.I):
            m = re.match(r'^(Given|When|Then|And|But)\s+(.+)$', line, re.I)
            assert m
            steps.append(_step_from_text(m.group(2), page, m.group(1)))
            if m.group(1).lower() == 'then':
                expected_parts.append(m.group(2))
            continue
        if section == 'preconditions':
            preconditions.append(_clean_step_text(line))
        elif section == 'expected':
            expected_parts.append(_clean_step_text(line))
        elif section == 'steps' or re.match(r'^\s*(?:step\s*)?\d+[\).:\-\s]+', line, re.I):
            steps.append(_step_from_text(line, page))
        elif not title:
            title = re.sub(r'^#+\s*', '', line).strip()
        else:
            steps.append(_step_from_text(line, page))
    if not steps and title:
        steps = [_step_from_text(title, page)]
    expected_result = ' '.join(x for x in expected_parts if x) or 'Expected business result should be verified.'
    if expected_parts and not any(step.get('action') == 'verify' for step in steps):
        steps.append(_step_from_text(expected_result, page, 'Then'))
    return {
        'id': _scenario_id(feature, index, scenario_id),
        'title': title or f'{feature.replace("_", " ").title()} test {index}',
        'feature': feature,
        'page': page,
        'priority': priority,
        'preconditions': [x for x in preconditions if x],
        'steps': [x for x in steps if x.get('target')],
        'expected_result': expected_result,
    }



def _jira_field(block: str, name: str) -> str:
    match = re.search(rf'^\s*{re.escape(name)}\s*:\s*(.*)$', block, re.I | re.M)
    return match.group(1).strip() if match else ''


def parse_jira_testcases(raw: str, feature: str) -> dict[str, Any]:
    blocks = [block.strip() for block in re.split(r'(?m)^\s*-{3,}\s*$', raw or '') if re.search(r'(?im)^\s*Jira Key\s*:', block)]
    parsed: list[dict[str, Any]] = []
    for block in blocks:
        key = _jira_field(block, 'Jira Key')
        issue_type = _jira_field(block, 'Issue Type') or 'Issue'
        title = _jira_field(block, 'Title') or f'Validate {key or feature}'
        priority = (_jira_field(block, 'Priority') or 'medium').lower()
        description_match = re.search(r'Description\s*/\s*Acceptance Criteria\s*:\s*(.*)$', block, re.I | re.S)
        description = description_match.group(1).strip() if description_match else ''
        page = pascal_case(feature) or 'Home'
        if re.search(r'(?im)^\s*(Given|When|Then)\s+', description):
            wrapped = f'Feature: {feature}\nScenario: {title}\n{description}'
            gherkin = parse_gherkin(wrapped, feature)
            scenario = (gherkin.get('scenarios') or [{}])[0]
            scenario.update({'id': key or scenario.get('id'), 'title': title, 'priority': priority, 'jira_issue_type': issue_type})
        else:
            lines = [line.strip(' \t-•') for line in description.splitlines() if line.strip()]
            steps = [_step_from_text(line, page) for line in lines if len(line) > 2]
            expected_candidates = [line for line in lines if any(word in line.lower() for word in ('should', 'expected', 'verify', 'displayed', 'visible', 'success', 'error'))]
            expected = expected_candidates[-1] if expected_candidates else f'{title} should satisfy the Jira acceptance criteria.'
            if not steps:
                steps = [_step_from_text(f'Validate {title}', page)]
            if not any(step.get('action') == 'verify' for step in steps):
                steps.append(_step_from_text(expected, page, 'Then'))
            scenario = {
                'id': key,
                'title': title,
                'feature': feature,
                'page': page,
                'priority': priority,
                'preconditions': [],
                'steps': steps,
                'expected_result': expected,
                'jira_issue_type': issue_type,
            }
        parsed.append(scenario)
    if len(parsed) > 1:
        non_epic = [scenario for scenario in parsed if str(scenario.get('jira_issue_type') or '').lower() != 'epic']
        if non_epic:
            parsed = non_epic
    for idx, scenario in enumerate(parsed, 1):
        scenario['id'] = _scenario_id(feature, idx, str(scenario.get('id') or ''))
        scenario.setdefault('feature', feature)
        scenario.setdefault('expected_result', f"{scenario.get('title')} should satisfy the Jira acceptance criteria.")
    return {
        'feature': feature,
        'source_type': 'jira_atlassian',
        'source_format': 'jira_issue_blocks',
        'scenario_count': len(parsed),
        'scenarios': parsed,
    }

def parse_plain_testcases(raw: str, feature: str) -> dict[str, Any]:
    if re.search(r'^\s*Feature\s*:', raw or '', re.I | re.M) or re.search(r'^\s*(Given|When|Then)\s+', raw or '', re.I | re.M):
        return parse_gherkin(raw, feature)
    scenarios = [_scenario_from_plain_block(block, feature, idx) for idx, block in enumerate(_split_plain_blocks(raw), 1)]
    scenarios = [s for s in scenarios if s.get('steps') or s.get('title')]
    return {
        'feature': feature,
        'source_type': 'module2_uploaded',
        'source_format': 'plain_steps',
        'scenario_count': len(scenarios),
        'scenarios': scenarios,
    }


def _extract_docx(uploaded_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(uploaded_bytes))
        lines: list[str] = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                lines.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    lines.append(' | '.join(values))
        return '\n'.join(lines)
    except Exception:
        with zipfile.ZipFile(io.BytesIO(uploaded_bytes)) as zf:
            xml = zf.read('word/document.xml').decode('utf-8', errors='replace')
        return '\n'.join(html.unescape(x) for x in re.findall(r'<w:t[^>]*>(.*?)</w:t>', xml, re.S))


def _extract_pdf(uploaded_bytes: bytes) -> str:
    import pypdf
    reader = pypdf.PdfReader(io.BytesIO(uploaded_bytes))
    return '\n'.join((page.extract_text() or '') for page in reader.pages)


def _find_header_index(headers: list[str], aliases: tuple[str, ...]) -> int | None:
    normalized = [re.sub(r'[^a-z0-9]+', ' ', h.lower()).strip() for h in headers]
    alias_values = [re.sub(r'[^a-z0-9]+', ' ', alias.lower()).strip() for alias in aliases]
    for alias_norm in alias_values:
        for idx, value in enumerate(normalized):
            if value == alias_norm:
                return idx
    for alias_norm in alias_values:
        if len(alias_norm) < 4:
            continue
        for idx, value in enumerate(normalized):
            if alias_norm in value:
                return idx
    return None


def _parse_excel(uploaded_bytes: bytes, feature: str) -> dict[str, Any]:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(uploaded_bytes), data_only=True, read_only=True)
    scenarios: list[dict[str, Any]] = []
    fallback_lines: list[str] = []
    for ws in wb.worksheets:
        rows = [[str(v).strip() if v is not None else '' for v in row] for row in ws.iter_rows(values_only=True)]
        rows = [r for r in rows if any(r)]
        if not rows:
            continue
        header_row = 0
        for idx, row in enumerate(rows[:15]):
            joined = ' '.join(row).lower()
            if any(key in joined for key in ('test case', 'testcase', 'expected', 'step', 'scenario')):
                header_row = idx
                break
        headers = rows[header_row]
        id_idx = _find_header_index(headers, ('test case id', 'testcase id', 'tc id', 'test case', 'test_case', 'tc', 'id'))
        title_idx = _find_header_index(headers, ('test case title', 'testcase title', 'summary', 'summery', 'scenario summary', 'scenario', 'title', 'test name'))
        step_idx = _find_header_index(headers, ('test step', 'step', 'steps', 'step description', 'action'))
        step_number_idx = _find_header_index(headers, ('step number', 'step no', 'step #', 'sequence'))
        test_data_idx = _find_header_index(headers, ('test data', 'testdata', 'input data', 'data', 'value'))
        expected_idx = _find_header_index(headers, ('expected result', 'expected outcome', 'expected'))
        pre_idx = _find_header_index(headers, ('precondition', 'prerequisite'))
        page_idx = _find_header_index(headers, ('page', 'screen', 'module'))
        priority_idx = _find_header_index(headers, ('priority',))
        if title_idx is None and step_idx is None:
            fallback_lines.append(f'# Sheet: {ws.title}')
            fallback_lines.extend(' | '.join(x for x in row if x) for row in rows)
            continue
        grouped: dict[str, dict[str, Any]] = {}
        last_key = ''
        for row_no, row in enumerate(rows[header_row + 1:], header_row + 2):
            def cell(index: int | None) -> str:
                return row[index].strip() if index is not None and index < len(row) else ''
            supplied_id = cell(id_idx)
            title = cell(title_idx)
            key = supplied_id or title or last_key or f'{ws.title}-{row_no}'
            last_key = key
            scenario = grouped.setdefault(key, {
                'id': supplied_id,
                'title': title or key,
                'feature': feature,
                'page': pascal_case(cell(page_idx)) if cell(page_idx) else 'Home',
                'priority': (cell(priority_idx) or 'medium').lower(),
                'preconditions': [],
                'steps': [],
                'expected_result': '',
                'source_sheet': ws.title,
            })
            if title and not scenario.get('title'):
                scenario['title'] = title
            if cell(pre_idx) and cell(pre_idx) not in scenario['preconditions']:
                scenario['preconditions'].append(cell(pre_idx))
            step_text = cell(step_idx)
            step_number = cell(step_number_idx)
            test_data = cell(test_data_idx)
            expected = cell(expected_idx)
            if step_text:
                step = _step_from_text(step_text, scenario['page'])
                step.update(_extract_test_data(step_text, test_data))
                if step_number:
                    step['step_number'] = step_number
                if expected:
                    step['expected'] = expected
                scenario['steps'].append(step)
            if expected:
                scenario['expected_result'] = expected
        for scenario in grouped.values():
            scenario['page'] = _infer_page_context(str(scenario.get('title') or ''), scenario.get('steps') or [], str(scenario.get('page') or 'Home'))
            for step in scenario.get('steps') or []:
                step['page'] = scenario['page']
                if step.get('value_sensitive') and str(scenario['page']).lower() == 'salesforce':
                    low_target = str(step.get('target') or '').lower()
                    if 'password' in low_target:
                        step['value_env'] = 'SALESFORCE_PASSWORD'
                    elif any(token in low_target for token in ('username', 'user name', 'email')):
                        step['value_env'] = 'SALESFORCE_USERNAME'
            if scenario.get('expected_result') and not any(step.get('action') == 'verify' for step in scenario.get('steps') or []):
                scenario['steps'].append(_step_from_text(scenario['expected_result'], scenario.get('page') or 'Home', 'Then'))
            if scenario.get('steps') or scenario.get('title'):
                scenarios.append(scenario)
    if not scenarios and fallback_lines:
        return parse_plain_testcases('\n'.join(fallback_lines), feature)
    for idx, scenario in enumerate(scenarios, 1):
        scenario['id'] = _scenario_id(feature, idx, scenario.get('id', ''))
        scenario['expected_result'] = scenario.get('expected_result') or 'Expected result from the spreadsheet should be verified.'
    return {
        'feature': feature,
        'source_type': 'module2_uploaded',
        'source_format': 'xlsx',
        'scenario_count': len(scenarios),
        'scenarios': scenarios,
    }


def extract_and_normalize_source(
    feature: str,
    pasted_json_or_steps: str = '',
    uploaded_bytes: bytes | None = None,
    uploaded_name: str = '',
    jira_story: str = '',
    jira_epic: str = '',
    source_mode: str = 'auto',
) -> dict[str, Any]:
    feature = _safe_feature(feature)
    uploaded_name = uploaded_name or ''
    suffix = Path(uploaded_name).suffix.lower()
    warnings: list[str] = []
    if uploaded_bytes and suffix in {'.xlsx', '.xlsm'}:
        payload = _parse_excel(uploaded_bytes, feature)
    else:
        extracted = ''
        if uploaded_bytes:
            if suffix in _TEXT_EXTENSIONS:
                extracted = uploaded_bytes.decode('utf-8', errors='replace')
            elif suffix == '.docx':
                extracted = _extract_docx(uploaded_bytes)
            elif suffix == '.pdf':
                extracted = _extract_pdf(uploaded_bytes)
            elif suffix == '.doc':
                try:
                    proc = subprocess.run(['antiword', '-'], input=uploaded_bytes, capture_output=True, timeout=30)
                    extracted = proc.stdout.decode('utf-8', errors='replace') if proc.returncode == 0 else ''
                except Exception:
                    extracted = ''
                if not extracted:
                    raise ValueError('Legacy .doc extraction requires antiword. Save the file as .docx, PDF, TXT, or MD and upload again.')
            else:
                extracted = uploaded_bytes.decode('utf-8', errors='replace')
        raw_parts: list[str] = []
        if jira_epic.strip():
            raw_parts.append('JIRA EPIC:\n' + jira_epic.strip())
        if jira_story.strip():
            raw_parts.append('JIRA STORY/TASK/BUG:\n' + jira_story.strip())
        if pasted_json_or_steps.strip():
            raw_parts.append(pasted_json_or_steps.strip())
        if extracted.strip():
            raw_parts.append(extracted.strip())
        raw = '\n\n---\n\n'.join(raw_parts)
        payload = _parse_json_payload(raw, feature)
        if payload is None:
            force_bdd = source_mode.lower() in {'bdd', 'gherkin', 'cucumber'} or suffix == '.feature'
            if source_mode.lower() == 'jira' or re.search(r'(?im)^\s*Jira Key\s*:', raw):
                payload = parse_jira_testcases(raw, feature)
            else:
                payload = parse_gherkin(raw, feature) if force_bdd else parse_plain_testcases(raw, feature)
    payload['feature'] = feature
    payload['scenario_count'] = len(payload.get('scenarios') or [])
    payload['source_file_name'] = uploaded_name
    payload['source_mode'] = source_mode
    payload['normalization_warnings'] = warnings
    if not payload['scenario_count']:
        raise ValueError('No executable testcase/scenario could be identified. Add Test Case/Scenario headings, numbered steps, or valid Given/When/Then content.')
    for idx, scenario in enumerate(payload['scenarios'], 1):
        scenario.setdefault('id', _scenario_id(feature, idx))
        scenario.setdefault('feature', feature)
        scenario.setdefault('page', 'Home')
        scenario.setdefault('priority', 'medium')
        scenario.setdefault('preconditions', [])
        scenario.setdefault('steps', [])
        scenario.setdefault('expected_result', 'Expected business result should be verified.')
    return payload


def save_normalized_source(payload: dict[str, Any]) -> dict[str, Any]:
    feature = _safe_feature(str(payload.get('feature') or 'feature'))
    path = feature_testcase_path('module2_uploaded', feature)
    path.parent.mkdir(parents=True, exist_ok=True)
    write_json(path, payload)
    write_functional_testcases_markdown(path, payload)
    write_active_context({
        'channel': 'module2_playwright_generator',
        'source_type': 'module2_uploaded',
        'requested_feature': feature,
        'parent_feature': feature,
        'features': [feature],
        'testcase_paths': [str(path.relative_to(REPO_ROOT))],
        'functional_testcases_reviewed': True,
        'review_gate': 'module2_uploaded_approved',
        'playwright_generated': False,
        'scenario_count': len(payload.get('scenarios') or []),
    })
    log_event('module2_testcases_load', f"Normalized {len(payload.get('scenarios') or [])} testcase(s) for existing-framework generation", status='done', progress=100, feature=feature)
    return {
        'ok': True,
        'testcase_file': str(path.relative_to(REPO_ROOT)),
        'markdown_file': str(path.with_name(path.name.replace('.scenarios.json', '.scenarios.md')).relative_to(REPO_ROOT)),
        'testcases': payload,
        'scenario_count': len(payload.get('scenarios') or []),
        'message': f"Identified and normalized {len(payload.get('scenarios') or [])} testcase(s). Each testcase will generate its own Playwright spec in the selected existing framework.",
    }


def _code_files(root: Path, dirs: list[str]) -> list[Path]:
    files: list[Path] = []
    for rel in dirs:
        base = root / rel
        if not base.exists():
            continue
        files.extend(p for p in base.rglob('*') if p.is_file() and p.suffix.lower() in _CODE_EXTENSIONS)
    return sorted(dict.fromkeys(files), key=lambda p: _rel(p, root).lower())


def _tokens(*values: str) -> list[str]:
    ignore = {
        'page', 'test', 'case', 'scenario', 'flow', 'validate', 'verify', 'the', 'and', 'with',
        'click', 'enter', 'select', 'open', 'navigate', 'displayed', 'visible', 'successfully',
        'button', 'field', 'section', 'record', 'created', 'expected', 'result', 'valid', 'new',
    }
    return [x for x in re.findall(r'[a-z0-9]+', ' '.join(values).lower()) if len(x) > 2 and x not in ignore]


def _file_score(path: Path, root: Path, tokens: list[str]) -> int:
    rel = _rel(path, root).lower()
    try:
        text = path.read_text(encoding='utf-8', errors='replace')[:120000].lower()
    except Exception:
        text = ''
    score = sum(6 for token in tokens if token in path.stem.lower())
    score += sum(2 for token in tokens if token in rel)
    score += sum(1 for token in tokens if token in text)
    if 'export class' in text:
        score += 3
    if 'page' in path.stem.lower():
        score += 1
    return score


def _is_concrete_page_file(path: Path) -> bool:
    name = path.stem.lower()
    if name in {'basepage', 'base_page', 'abstractpage', 'pagebase'} or name.startswith(('basepage.', 'abstractpage.')):
        return False
    if any(part.lower() in {'utils', 'utilities', 'helpers', 'fixtures', 'config', 'configs'} for part in path.parts):
        return False
    try:
        text = path.read_text(encoding='utf-8', errors='replace')[:8000]
    except Exception:
        return False
    if re.search(r'export\s+abstract\s+class\s+', text):
        return False
    return bool(re.search(r'export\s+(?:default\s+)?class\s+', text))


def _rank_files(root: Path, files: list[Path], tokens: list[str]) -> list[dict[str, Any]]:
    ranked = [{'path': _rel(path, root), 'score': _file_score(path, root, tokens)} for path in files]
    return sorted(ranked, key=lambda x: (-x['score'], len(x['path']), x['path'].lower()))


def _choose_test_dir(root: Path, profile: dict[str, Any], feature: str, explicit: str = '') -> tuple[Path, list[str]]:
    selected = _resolve_under_root(root, explicit) if explicit else None
    reasons: list[str] = []
    if selected:
        reasons.append('User explicitly selected this test folder.')
        return selected, reasons
    configured = profile.get('configured_test_dirs') or []
    if configured:
        configured_dir = root / configured[0]
        generated_subdir = configured_dir / 'generated'
        if generated_subdir.exists() and (any(generated_subdir.glob('*.spec.*')) or (generated_subdir / '.gitkeep').exists()):
            reasons.append('Selected the existing generated-spec subfolder under Playwright testDir to preserve framework conventions and execution scripts.')
            return generated_subdir, reasons
        reasons.append('Selected from Playwright config testDir discovered by framework learning.')
        return configured_dir, reasons
    roots = profile.get('discovered_test_roots') or []
    if roots:
        reasons.append('Selected from recursively proven executable Playwright test root.')
        return root / roots[0], reasons
    specs = profile.get('executable_specs') or []
    if specs:
        reasons.append('Selected from the parent folder of an existing executable spec.')
        return (root / specs[0]).parent, reasons
    reasons.append('No test root exists; tests is the conventional fallback and will be created inside the selected framework.')
    return root / 'tests', reasons


def _class_name(path: Path, fallback: str) -> str:
    if path.exists():
        text = path.read_text(encoding='utf-8', errors='replace')
        match = re.search(r'export\s+(?:default\s+)?class\s+([A-Za-z_$][\w$]*)', text)
        if match:
            return match.group(1)
    return fallback


def _placement_for_scenario(root: Path, profile: dict[str, Any], scenario: dict[str, Any], feature: str, explicit_page: str = '', explicit_locator: str = '') -> dict[str, Any]:
    components = profile.get('component_directory_model') or {}
    page_dirs = components.get('page_dirs') or []
    object_dirs = components.get('page_object_dirs') or []
    page_files = [path for path in _code_files(root, page_dirs) if _is_concrete_page_file(path)]
    object_files = _code_files(root, object_dirs)
    step_context = ' '.join(str(step.get('target') or '') for step in scenario.get('steps') or [])
    tokens = _tokens(feature, str(scenario.get('page') or ''), str(scenario.get('title') or ''), step_context)
    explicit_page_path = _resolve_under_root(root, explicit_page, allow_missing=True) if explicit_page else None
    explicit_locator_path = _resolve_under_root(root, explicit_locator, allow_missing=True) if explicit_locator else None
    ranked_pages = _rank_files(root, page_files, tokens)
    ranked_objects = _rank_files(root, object_files, tokens)
    page_path = explicit_page_path
    locator_path = explicit_locator_path
    page_reason = 'explicit_user_selection' if page_path else ''
    locator_reason = 'explicit_user_selection' if locator_path else ''
    ambiguous = False
    if page_path is None and ranked_pages:
        context_tokens = _tokens(str(scenario.get('page') or ''))
        matching_context = next((item for item in ranked_pages if any(token in Path(item['path']).stem.lower() for token in context_tokens)), None)
        selected = matching_context or ranked_pages[0]
        selected_name = Path(selected['path']).stem.lower()
        context_matches_name = any(token in selected_name for token in context_tokens)
        if not context_tokens or str(scenario.get('page') or '').lower() in {'home', 'generated'} or context_matches_name:
            page_path = root / selected['path']
            page_reason = 'best_existing_page_match'
            if matching_context is None and len(ranked_pages) > 1 and ranked_pages[0]['score'] == ranked_pages[1]['score'] and ranked_pages[0]['score'] > 0:
                ambiguous = True
        else:
            page_reason = 'no_existing_page_matches_application_context'
    if locator_path is None and page_path is not None:
        linked = next((candidate for candidate in object_files if _locator_binding(page_path, candidate) is not None), None)
        if linked is not None:
            locator_path = linked
            locator_reason = 'existing_locator_repository_imported_or_instantiated_by_selected_page'
        else:
            locator_path = page_path
            locator_reason = 'safe_unified_page_object_placement; no linked locator repository was detected'
    return {
        'scenario_id': scenario.get('id'),
        'scenario_title': scenario.get('title'),
        'recommended_page_file': _rel(page_path, root) if page_path else '',
        'recommended_locator_file': _rel(locator_path, root) if locator_path else '',
        'page_reason': page_reason or 'no_existing_page_match',
        'locator_reason': locator_reason or 'no_existing_locator_match',
        'ambiguous': ambiguous,
        'page_candidates': ranked_pages[:12],
        'locator_candidates': ranked_objects[:12],
    }


def preview_generation_placement(
    framework_path: str,
    feature: str,
    target_test_folder: str = '',
    target_page_file: str = '',
    target_locator_file: str = '',
    placement_mode: str = 'confirm_if_ambiguous',
) -> dict[str, Any]:
    root = Path(framework_path).expanduser().resolve()
    if not root.exists():
        return {'ok': False, 'error': f'Framework path does not exist: {root}'}
    feature = _safe_feature(feature)
    testcase_path = feature_testcase_path('module2_uploaded', feature)
    if not testcase_path.exists():
        return {'ok': False, 'error': 'Load/normalize testcase source first.'}
    payload = read_json(testcase_path)
    profile = build_structure_profile(root, limit=7000)
    test_dir, test_reasons = _choose_test_dir(root, profile, feature, target_test_folder)
    placements = [
        _placement_for_scenario(root, profile, scenario, feature, target_page_file, target_locator_file)
        for scenario in payload.get('scenarios') or []
    ]
    unresolved = [p for p in placements if not p.get('recommended_page_file')]
    ambiguous = [p for p in placements if p.get('ambiguous')]
    needs_confirmation = bool(unresolved or (ambiguous and placement_mode == 'confirm_if_ambiguous'))
    return {
        'ok': True,
        'framework_path': str(root),
        'feature': feature,
        'scenario_count': len(placements),
        'recommended_test_folder': _rel(test_dir, root),
        'test_folder_reasons': test_reasons,
        'placements': placements,
        'needs_user_confirmation': needs_confirmation,
        'unresolved_scenarios': [p['scenario_id'] for p in unresolved],
        'ambiguous_scenarios': [p['scenario_id'] for p in ambiguous],
        'message': 'Placement preview completed. No framework source file was modified.' if not needs_confirmation else 'Placement is ambiguous or no reusable page file exists. Select the page/locator target or allow a new support file before generation.',
    }


def _backup_file(root: Path, path: Path, backup_root: Path) -> str:
    rel = _rel(path, root)
    if path.exists():
        target = backup_root / rel
        if not target.exists():
            target.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(path, target)
    return rel


def _rollback_generation(root: Path, backup_root: Path, created_paths: set[str]) -> dict[str, Any]:
    restored: list[str] = []
    deleted: list[str] = []
    if backup_root.exists():
        for backup in sorted((p for p in backup_root.rglob('*') if p.is_file()), key=lambda p: len(p.parts)):
            rel = backup.relative_to(backup_root)
            target = root / rel
            target.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(backup, target)
            restored.append(rel.as_posix())
    for rel in sorted(created_paths, key=lambda value: len(Path(value).parts), reverse=True):
        target = root / rel
        if target.exists() and target.is_file():
            target.unlink()
            deleted.append(rel)
        parent = target.parent
        while parent != root and parent.exists():
            try:
                parent.rmdir()
            except OSError:
                break
            parent = parent.parent
    return {
        'performed': True,
        'restored_files': sorted(dict.fromkeys(restored)),
        'deleted_created_files': sorted(dict.fromkeys(deleted)),
        'reason': 'Playwright validation failed, so AstraHeal restored every pre-existing file from backup and removed every newly created source/spec file.',
    }


def _matching_brace_index(text: str, open_index: int) -> int:
    depth = 0
    quote = ''
    escaped = False
    line_comment = False
    block_comment = False
    i = open_index
    while i < len(text):
        char = text[i]
        nxt = text[i + 1] if i + 1 < len(text) else ''
        if line_comment:
            if char == '\n':
                line_comment = False
            i += 1
            continue
        if block_comment:
            if char == '*' and nxt == '/':
                block_comment = False
                i += 2
                continue
            i += 1
            continue
        if quote:
            if escaped:
                escaped = False
            elif char == '\\':
                escaped = True
            elif char == quote:
                quote = ''
            i += 1
            continue
        if char == '/' and nxt == '/':
            line_comment = True
            i += 2
            continue
        if char == '/' and nxt == '*':
            block_comment = True
            i += 2
            continue
        if char in {'\'', '"', '`'}:
            quote = char
            i += 1
            continue
        if char == '{':
            depth += 1
        elif char == '}':
            depth -= 1
            if depth == 0:
                return i
        i += 1
    return -1


def _class_close_index(text: str, class_name: str = '') -> int:
    name = rf'\s+{re.escape(class_name)}\b' if class_name else r'\s+[A-Za-z_$][\w$]*\b'
    match = re.search(rf'export\s+(?:default\s+)?(?:abstract\s+)?class{name}[^{{]*{{', text)
    if not match:
        return -1
    open_index = text.find('{', match.start())
    return _matching_brace_index(text, open_index)


def _append_member(path: Path, member: str, symbol: str) -> bool:
    text = path.read_text(encoding='utf-8', errors='replace') if path.exists() else ''
    if re.search(rf'\b{re.escape(symbol)}\b', text):
        return False
    class_name = _class_name(path, '')
    idx = _class_close_index(text, class_name)
    if idx < 0:
        raise ValueError(f'Could not safely identify exported class boundary in {path}')
    path.write_text(text[:idx].rstrip() + '\n\n' + member.rstrip() + '\n' + text[idx:], encoding='utf-8')
    return True


def _locator_name(target: str) -> str:
    base = camel_case(re.sub(r'[^A-Za-z0-9 ]+', ' ', _semantic_target(target))) or 'element'
    return base if base.lower().endswith('locator') else base + 'Locator'


def _method_name(action: str, target: str) -> str:
    words = re.findall(r'[A-Za-z0-9]+', f'{action} {_semantic_target(target)}')[:7]
    name = ''.join(x[:1].upper() + x[1:] for x in words) or 'PerformStep'
    lower = name[:1].lower() + name[1:]
    return lower if lower.startswith(('click', 'fill', 'verify', 'select', 'goto', 'navigate', 'perform')) else 'perform' + name


def _locator_expression(target: str, action: str) -> str:
    label = re.sub(r"['`\n\r]", '', _semantic_target(target))[:120]
    low = f'{action} {target}'.lower()
    escaped = re.escape(label).replace('/', r'\/')
    if any(x in low for x in ('email', 'username', 'password', 'textbox', 'input', 'field')):
        return f"this.page.getByRole('textbox', {{ name: /{escaped}/i }})"
    if any(x in low for x in ('button', 'click', 'submit', 'save', 'continue', 'login', 'sign in')):
        return f"this.page.getByRole('button', {{ name: /{escaped}/i }})"
    if any(x in low for x in ('link', 'navigate', 'menu')):
        return f"this.page.getByRole('link', {{ name: /{escaped}/i }})"
    return f"this.page.getByText(/{escaped}/i)"


def _ts_single_quote(value: str) -> str:
    return str(value).replace('\\', '\\\\').replace("'", "\\'").replace('\r', ' ').replace('\n', ' ')


def _locator_definition(target: str, action: str) -> str:
    label = _ts_single_quote(_semantic_target(target)[:120])
    low = f'{action} {target}'.lower()
    if any(x in low for x in ('email', 'username', 'password', 'textbox', 'input', 'field', 'first name', 'last name', 'mobile', 'notes', 'search box')):
        return f"{{ strategy: 'role', role: 'textbox', value: '{label}', description: '{label}', fallbacks: [{{ strategy: 'label', value: '{label}' }}, {{ strategy: 'placeholder', value: '{label}' }}] }}"
    if any(x in low for x in ('button', 'click', 'submit', 'save', 'continue', 'login', 'sign in', 'icon')):
        return f"{{ strategy: 'role', role: 'button', value: '{label}', description: '{label}', fallbacks: [{{ strategy: 'text', value: '{label}' }}] }}"
    if any(x in low for x in ('link', 'navigate', 'menu', 'tab')):
        return f"{{ strategy: 'role', role: 'link', value: '{label}', description: '{label}', fallbacks: [{{ strategy: 'text', value: '{label}' }}] }}"
    return f"{{ strategy: 'text', value: '{label}', description: '{label}' }}"


def _create_page_file(path: Path, class_name: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        "import { Page } from '@playwright/test';\n\n"
        f"export class {class_name} {{\n"
        "  constructor(private readonly page: Page) {}\n"
        "}\n",
        encoding='utf-8',
    )


def _create_framework_support_files(root: Path, profile: dict[str, Any], class_name: str) -> tuple[Path, Path, list[Path]]:
    components = profile.get('component_directory_model') or {}
    page_dir = root / ((components.get('page_dirs') or ['pages'])[0])
    object_dirs = components.get('page_object_dirs') or []
    page_path = page_dir / f'{class_name}.ts'
    created: list[Path] = []
    base_candidates = [path for path in _code_files(root, components.get('page_dirs') or ['pages']) if path.stem.lower() in {'basepage', 'base_page'}]
    locator_factory = next((path for path in root.rglob('locatorFactory.ts') if path.is_file()), None)
    if base_candidates and object_dirs and locator_factory:
        base_path = base_candidates[0]
        object_dir = root / object_dirs[0]
        locator_name = f'{class_name}Objects'
        locator_path = object_dir / f'{class_name}.objects.ts'
        if not locator_path.exists():
            locator_path.parent.mkdir(parents=True, exist_ok=True)
            factory_import = _relative_import(locator_path, locator_factory)
            locator_path.write_text(
                f"import type {{ LocatorDefinition }} from '{factory_import}';\n\n"
                f"export const {locator_name} = {{\n"
                f"}} satisfies Record<string, LocatorDefinition>;\n",
                encoding='utf-8',
            )
            created.append(locator_path)
        if not page_path.exists():
            page_path.parent.mkdir(parents=True, exist_ok=True)
            base_import = _relative_import(page_path, base_path)
            locator_import = _relative_import(page_path, locator_path)
            page_path.write_text(
                "import type { Page } from '@playwright/test';\n"
                f"import {{ BasePage }} from '{base_import}';\n"
                f"import {{ {locator_name} }} from '{locator_import}';\n\n"
                f"export class {class_name} extends BasePage {{\n"
                "  constructor(page: Page) {\n"
                "    super(page);\n"
                "  }\n"
                "}\n",
                encoding='utf-8',
            )
            created.append(page_path)
        return page_path, locator_path, created
    if not page_path.exists():
        _create_page_file(page_path, class_name)
        created.append(page_path)
    return page_path, page_path, created


def _object_literal_name(path: Path) -> str:
    if not path.exists():
        return ''
    text = path.read_text(encoding='utf-8', errors='replace')
    match = re.search(r'export\s+const\s+([A-Za-z_$][\w$]*)\s*=\s*{', text)
    return match.group(1) if match else ''


def _locator_binding(page_path: Path, locator_path: Path) -> dict[str, str] | None:
    if page_path.resolve() == locator_path.resolve():
        return {'style': 'direct_class_member', 'locator_expression': 'this.{locator}', 'reference': 'this'}
    if not page_path.exists() or not locator_path.exists():
        return None
    page_text = page_path.read_text(encoding='utf-8', errors='replace')
    object_name = _object_literal_name(locator_path)
    if object_name and re.search(rf'\b{re.escape(object_name)}\b', page_text) and ('getLocator' in page_text or re.search(r'extends\s+BasePage\b', page_text)):
        return {
            'style': 'locator_definition_object',
            'collection': object_name,
            'locator_expression': f'this.getLocator({object_name}.{{locator}})',
            'reference': object_name,
        }
    locator_class = _class_name(locator_path, pascal_case(locator_path.stem))
    if locator_class not in page_text:
        return None
    patterns = [
        rf'this\.([A-Za-z_$][\w$]*)\s*=\s*new\s+{re.escape(locator_class)}\s*\(',
        rf'(?:private|protected|public|readonly|private\s+readonly|protected\s+readonly)?\s*([A-Za-z_$][\w$]*)\s*[:=][^;\n]*\b{re.escape(locator_class)}\b',
    ]
    for pattern in patterns:
        match = re.search(pattern, page_text)
        if match:
            reference = f'this.{match.group(1)}'
            return {'style': 'locator_class_instance', 'locator_expression': reference + '.{locator}', 'reference': reference}
    return None


def _append_locator(path: Path, locator: str, target: str, action: str, binding: dict[str, str]) -> bool:
    text = path.read_text(encoding='utf-8', errors='replace') if path.exists() else ''
    if re.search(rf'\b{re.escape(locator)}\b', text):
        return False
    if binding.get('style') == 'locator_definition_object':
        object_name = binding.get('collection') or _object_literal_name(path)
        match = re.search(rf'export\s+const\s+{re.escape(object_name)}\s*=\s*{{', text)
        if not match:
            raise ValueError(f'Could not safely identify locator object boundary in {path}')
        open_index = text.find('{', match.start())
        close_index = _matching_brace_index(text, open_index)
        if close_index < 0:
            raise ValueError(f'Could not safely identify locator object closing brace in {path}')
        member = f"  {locator}: {_locator_definition(target, action)},"
        path.write_text(text[:close_index].rstrip() + '\n' + member + '\n' + text[close_index:], encoding='utf-8')
        return True
    member = f"  readonly {locator} = {_locator_expression(target, action)};"
    return _append_member(path, member, locator)



def _linked_locator_reference(page_path: Path, locator_path: Path) -> str | None:
    binding = _locator_binding(page_path, locator_path)
    return binding.get('reference') if binding else None

def _relative_import(from_file: Path, to_file: Path) -> str:
    rel = os.path.relpath(to_file.with_suffix(''), from_file.parent).replace('\\', '/')
    return rel if rel.startswith('.') else './' + rel


def _unique_spec_path(test_dir: Path, feature: str, scenario: dict[str, Any], index: int) -> Path:
    label = safe_id(str(scenario.get('id') or scenario.get('title') or index)).replace('_', '-').lower()
    base = f'{feature}-{label}.spec.ts'
    candidate = test_dir / base
    counter = 2
    while candidate.exists():
        candidate = test_dir / f'{feature}-{label}-{counter}.spec.ts'
        counter += 1
    return candidate


def _step_call(step: dict[str, Any], method: str) -> str:
    action = str(step.get('action') or '').lower()
    value = str(step.get('value') or '').strip()
    value_env = re.sub(r'[^A-Z0-9_]+', '_', str(step.get('value_env') or '').upper()).strip('_')
    expected = str(step.get('expected') or '').strip()
    if value_env:
        arg_expression = f"process.env.{value_env} || {json.dumps('${' + value_env + '_REQUIRED}')}"
    elif value:
        arg_expression = json.dumps(value)
    else:
        fallback_env = _environment_name_for_step(str(step.get('target') or 'test data'))
        arg_expression = f"process.env.{fallback_env} || {json.dumps('${' + fallback_env + '_REQUIRED}')}"
    if action in {'goto', 'open', 'launch', 'navigate'}:
        url_match = re.search(r'https?://[^\s]+', str(step.get('target') or ''))
        if value:
            arg = json.dumps(value)
        elif url_match:
            arg = json.dumps(url_match.group(0).rstrip('.,;'))
        else:
            arg = "process.env.BASE_URL || process.env.TEST_BASE_URL || '/'"
        return f'await screen.{method}({arg});'
    if action in {'fill', 'type', 'enter', 'select'}:
        return f'await screen.{method}({arg_expression});'
    if action in {'verify', 'assert', 'expect', 'validate'} and expected:
        return f'await screen.{method}({json.dumps(expected)});'
    return f'await screen.{method}();'


def _method_member(action: str, method: str, locator_ref: str) -> str:
    if action in {'fill', 'type', 'enter'}:
        return f"  async {method}(value: string) {{\n    await {locator_ref}.fill(value);\n  }}"
    if action == 'select':
        return f"  async {method}(value: string) {{\n    const target = {locator_ref};\n    await target.selectOption({{ label: value }}).catch(async () => {{\n      await target.click();\n      await this.page.getByRole('option', {{ name: value }}).click();\n    }});\n  }}"
    if action in {'verify', 'assert', 'expect', 'validate'}:
        return f"  async {method}(expectedText?: string) {{\n    await {locator_ref}.waitFor({{ state: 'visible' }});\n    if (expectedText) {{\n      const actualText = (await {locator_ref}.textContent()) || '';\n      if (!actualText.includes(expectedText)) throw new Error(`Expected text not found: ${{expectedText}}`);\n    }}\n  }}"
    return f"  async {method}() {{\n    await {locator_ref}.click();\n  }}"


def _generate_report(root: Path, plan: dict[str, Any]) -> tuple[Path, Path]:
    report_dir = root / '.aiqa-history' / 'add-new-tests'
    report_dir.mkdir(parents=True, exist_ok=True)
    json_path = report_dir / f"{plan['feature']}-generation-report.json"
    html_path = report_dir / f"{plan['feature']}-generation-report.html"
    write_json(json_path, plan)
    rows = ''.join(
        f"<tr><td>{html.escape(str(x.get('scenario_id')))}</td><td>{html.escape(str(x.get('scenario_title')))}</td><td><code>{html.escape(str(x.get('spec_file')))}</code></td><td>{html.escape(str(x.get('page_file')))}</td></tr>"
        for x in plan.get('scenario_outputs') or []
    )
    html_path.write_text(
        "<!doctype html><html><head><meta charset='utf-8'><title>Add New Tests Report</title>"
        "<style>body{font-family:Segoe UI,Arial;margin:24px;background:#f8fafc;color:#0f172a}.card{background:white;border:1px solid #dbe3ef;border-radius:12px;padding:16px;margin:14px 0}table{border-collapse:collapse;width:100%}td,th{border-bottom:1px solid #e2e8f0;padding:8px;text-align:left}code,pre{background:#0f172a;color:#dbeafe;padding:3px 6px;border-radius:6px;white-space:pre-wrap}</style></head><body>"
        f"<h1>Add New Tests Generation Report</h1><div class='card'><b>Framework:</b> <code>{html.escape(str(root))}</code><br/><b>Feature:</b> {html.escape(str(plan.get('feature')))}<br/><b>Scenarios:</b> {plan.get('scenario_count')}<br/><b>Files actually changed:</b> {len(plan.get('changed_files') or [])}</div>"
        f"<div class='card'><h2>Scenario-to-script mapping</h2><table><tr><th>Scenario</th><th>Title</th><th>Generated spec</th><th>Reused/updated page file</th></tr>{rows}</table></div>"
        f"<div class='card'><h2>Created and reused symbols</h2><pre>{html.escape(json.dumps({'created': plan.get('created_symbols'), 'reused': plan.get('reused_symbols')}, indent=2, ensure_ascii=False))}</pre></div>"
        f"<div class='card'><h2>Validation</h2><pre>{html.escape(json.dumps(plan.get('validation'), indent=2, ensure_ascii=False))}</pre></div></body></html>",
        encoding='utf-8',
    )
    return json_path, html_path


def _validate_typescript_syntax(root: Path, source_files: list[str]) -> dict[str, Any]:
    node = shutil.which('node')
    files = [rel for rel in source_files if Path(rel).suffix.lower() in {'.ts', '.tsx', '.js', '.jsx', '.mjs', '.cjs'} and (root / rel).exists()]
    if not node or not files:
        return {'ok': None, 'skipped': True, 'reason': 'Node.js or generated TypeScript/JavaScript source files are unavailable.'}
    script = r"""
const fs = require('fs');
let ts;
try { ts = require('typescript'); }
catch (error) { console.log(JSON.stringify({ skipped: true, reason: String(error) })); process.exit(0); }
const diagnostics = [];
for (const file of process.argv.slice(1)) {
  const text = fs.readFileSync(file, 'utf8');
  const kind = file.endsWith('.tsx') ? ts.ScriptKind.TSX : file.endsWith('.jsx') ? ts.ScriptKind.JSX : file.endsWith('.js') ? ts.ScriptKind.JS : ts.ScriptKind.TS;
  const source = ts.createSourceFile(file, text, ts.ScriptTarget.Latest, true, kind);
  for (const item of source.parseDiagnostics || []) {
    const pos = source.getLineAndCharacterOfPosition(item.start || 0);
    diagnostics.push({ file, line: pos.line + 1, column: pos.character + 1, code: item.code, message: ts.flattenDiagnosticMessageText(item.messageText, ' ') });
  }
}
console.log(JSON.stringify({ skipped: false, diagnostics }));
process.exit(diagnostics.length ? 2 : 0);
"""
    try:
        env = os.environ.copy()
        local_typescript = root / 'node_modules' / 'typescript'
        if not local_typescript.exists():
            npm = shutil.which('npm')
            if npm:
                npm_root = subprocess.run([npm, 'root', '-g'], capture_output=True, text=True, timeout=15)
                global_root = (npm_root.stdout or '').strip() if npm_root.returncode == 0 else ''
                if global_root:
                    existing_node_path = env.get('NODE_PATH', '').strip()
                    env['NODE_PATH'] = os.pathsep.join(x for x in (global_root, existing_node_path) if x)
        proc = subprocess.run([node, '-e', script, *files], cwd=root, capture_output=True, text=True, timeout=60, env=env)
        raw = (proc.stdout or '').strip().splitlines()
        payload = json.loads(raw[-1]) if raw else {}
        if payload.get('skipped'):
            return {'ok': None, 'skipped': True, 'reason': payload.get('reason') or 'TypeScript package is unavailable.'}
        diagnostics = payload.get('diagnostics') or []
        return {
            'ok': not diagnostics,
            'skipped': False,
            'diagnostic_count': len(diagnostics),
            'diagnostics': diagnostics[:100],
            'files_checked': files,
        }
    except Exception as exc:
        return {'ok': None, 'skipped': True, 'reason': f'{type(exc).__name__}: {exc}'}


def _playwright_command(root: Path) -> list[str] | None:
    local = root / 'node_modules' / '.bin' / ('playwright.cmd' if os.name == 'nt' else 'playwright')
    if local.exists():
        return [str(local)]
    npx = shutil.which('npx')
    return [npx, '--no-install', 'playwright'] if npx else None


def _validate_specs(root: Path, specs: list[str], enabled: bool, source_files: list[str] | None = None) -> dict[str, Any]:
    if not enabled:
        return {'ok': None, 'skipped': True, 'reason': 'Validation was disabled by the user.'}
    syntax = _validate_typescript_syntax(root, source_files or specs)
    if syntax.get('ok') is False:
        return {
            'ok': False,
            'skipped': False,
            'stage': 'typescript_syntax',
            'plain_english_reason': 'Generated TypeScript/JavaScript contains a syntax error. Playwright execution was not attempted, and rollback is required.',
            'syntax': syntax,
        }
    command = _playwright_command(root)
    if not command or not (root / 'package.json').exists():
        return {
            'ok': None,
            'skipped': True,
            'reason': 'Local Playwright CLI or package.json is unavailable; static syntax validation completed but Playwright --list could not run.',
            'syntax': syntax,
        }
    args = [*command, 'test', *specs, '--list']
    try:
        proc = subprocess.run(args, cwd=root, capture_output=True, text=True, timeout=180)
        stderr = proc.stderr[-12000:]
        stdout = proc.stdout[-12000:]
        return {
            'ok': proc.returncode == 0,
            'skipped': False,
            'stage': 'playwright_list',
            'command': ' '.join(args),
            'returncode': proc.returncode,
            'stdout': stdout,
            'stderr': stderr,
            'syntax': syntax,
            'plain_english_reason': ('Playwright successfully discovered every generated testcase.' if proc.returncode == 0 else 'Playwright could not load one or more generated specs or their dependent page/locator files. Review the first error and impacted file below; AstraHeal will roll back the attempted changes.'),
        }
    except Exception as exc:
        return {'ok': False, 'skipped': False, 'error': f'{type(exc).__name__}: {exc}'}


def generate_existing_framework_tests(
    framework_path: str,
    feature: str,
    provider: str = 'deterministic',
    model: str = 'llama3',
    base_url: str = '',
    target_test_folder: str = '',
    target_page_file: str = '',
    target_locator_file: str = '',
    placement_mode: str = 'confirm_if_ambiguous',
    allow_new_support_files: bool = True,
    validate_generated: bool = True,
    bdd_output_mode: str = 'playwright_specs',
) -> dict[str, Any]:
    root = Path(framework_path).expanduser().resolve()
    if not root.exists():
        return {'ok': False, 'error': f'Framework path does not exist: {root}'}
    feature = _safe_feature(feature)
    testcase_path = feature_testcase_path('module2_uploaded', feature)
    if not testcase_path.exists():
        return {'ok': False, 'error': 'Load/normalize testcase source first.'}
    payload = read_json(testcase_path)
    scenarios = payload.get('scenarios') or []
    if not scenarios:
        return {'ok': False, 'error': 'No normalized scenarios are available.'}
    log_event('module2_existing_framework', f'Preparing {len(scenarios)} new Playwright specs in the selected framework', progress=5, feature=feature)
    profile = build_structure_profile(root, limit=7000)
    try:
        deep = build_deep_framework_understanding(root, base_url=base_url)
    except Exception as exc:
        deep = {'ok': False, 'warning': f'{type(exc).__name__}: {exc}'}
    try:
        from qa_pipeline.mcp.playwright_mcp import mcp_status
        mcp_readiness = mcp_status(headless=False, probe_server=False)
    except Exception as exc:
        mcp_readiness = {'mcp_probe_ok': False, 'error': f'{type(exc).__name__}: {exc}'}
    preview = preview_generation_placement(str(root), feature, target_test_folder, target_page_file, target_locator_file, placement_mode)
    if not preview.get('ok'):
        return preview
    if placement_mode == 'confirm_if_ambiguous' and not target_page_file:
        ambiguous = preview.get('ambiguous_scenarios') or []
        unresolved = preview.get('unresolved_scenarios') or []
        if ambiguous or (unresolved and not allow_new_support_files):
            return {
                'ok': False,
                'needs_user_input': True,
                'placement_preview': preview,
                'message': 'Generation stopped safely before changing files because page/locator placement is ambiguous or support-file creation was not permitted. Select the target page/locator file, allow required support-file creation, or choose automatic placement explicitly.',
            }
    for item in preview.get('placements') or []:
        page_rel = item.get('recommended_page_file') or ''
        locator_rel = item.get('recommended_locator_file') or ''
        if page_rel and locator_rel and page_rel != locator_rel:
            page_candidate = root / page_rel
            locator_candidate = root / locator_rel
            if _locator_binding(page_candidate, locator_candidate) is None:
                return {
                    'ok': False,
                    'needs_user_input': True,
                    'placement_preview': preview,
                    'message': f"The selected locator file {locator_rel} is not visibly linked to {page_rel}. Select the same page file for locators, or choose an existing locator repository already imported and instantiated by that page class. No files were changed.",
                }
    test_dir = root / preview['recommended_test_folder']
    test_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    backup_root = root / '.aiqa-history' / 'backups' / 'add-new-tests' / timestamp
    changed_files: list[str] = []
    created_paths: set[str] = set()
    created_symbols: list[dict[str, Any]] = []
    reused_symbols: list[dict[str, Any]] = []
    scenario_outputs: list[dict[str, Any]] = []
    placement_by_id = {p['scenario_id']: p for p in preview.get('placements') or []}

    for index, scenario in enumerate(scenarios, 1):
        placement = placement_by_id.get(scenario.get('id')) or {}
        page_rel = placement.get('recommended_page_file')
        page_path = root / page_rel if page_rel else None
        auto_locator_path: Path | None = None
        if page_path is None or not page_path.exists():
            if not allow_new_support_files:
                return {'ok': False, 'needs_user_input': True, 'placement_preview': preview, 'message': f"No reusable page file exists for {scenario.get('id')}. Select a page file or allow creation of a support file only when required."}
            class_name = pascal_case(str(scenario.get('page') or feature)) or 'GeneratedPage'
            if not class_name.endswith('Page'):
                class_name += 'Page'
            page_path, auto_locator_path, created_support = _create_framework_support_files(root, profile, class_name)
            for created_path in created_support:
                created_paths.add(_rel(created_path, root))
                changed_files.append(_rel(created_path, root))
                created_symbols.append({'file': _rel(created_path, root), 'kind': 'support_file', 'name': created_path.stem, 'reason': 'No semantically matching concrete page file existed; one reusable application page/locator pair was created and shared by all related testcases.'})
        class_name = _class_name(page_path, pascal_case(page_path.stem))
        locator_rel = placement.get('recommended_locator_file') or (_rel(auto_locator_path, root) if auto_locator_path else _rel(page_path, root))
        locator_path = root / locator_rel
        binding = _locator_binding(page_path, locator_path)
        if binding is None:
            return {'ok': False, 'needs_user_input': True, 'placement_preview': preview, 'message': f'Locator file {locator_rel} is not linked to page file {_rel(page_path, root)}. No further files were changed.'}
        _backup_file(root, page_path, backup_root)
        if locator_path.resolve() != page_path.resolve():
            _backup_file(root, locator_path, backup_root)
        text = page_path.read_text(encoding='utf-8', errors='replace')
        locator_text = locator_path.read_text(encoding='utf-8', errors='replace')
        calls: list[str] = []
        for step in scenario.get('steps') or []:
            action = str(step.get('action') or 'perform').lower()
            target = str(step.get('target') or 'element')
            if action in {'goto', 'open', 'launch', 'navigate'}:
                method = 'gotoApplication'
                if re.search(r'\bgotoApplication\s*\(', text):
                    reused_symbols.append({'file': _rel(page_path, root), 'kind': 'method', 'name': method, 'scenario_id': scenario.get('id')})
                else:
                    member = ("  async gotoApplication(url: string) {\n    await this.goto(url);\n  }" if re.search(r'extends\s+BasePage\b', text) else "  async gotoApplication(url: string) {\n    await this.page.goto(url);\n  }")
                    if _append_member(page_path, member, method):
                        text += '\n' + member
                        changed_files.append(_rel(page_path, root))
                        created_symbols.append({'file': _rel(page_path, root), 'kind': 'method', 'name': method, 'scenario_id': scenario.get('id')})
                calls.append(_step_call(step, method))
                continue
            locator = _locator_name(target)
            method = _method_name(action, target)
            if re.search(rf'\b{re.escape(locator)}\b', locator_text):
                reused_symbols.append({'file': _rel(locator_path, root), 'kind': 'locator', 'name': locator, 'scenario_id': scenario.get('id')})
            else:
                if _append_locator(locator_path, locator, target, action, binding):
                    locator_text = locator_path.read_text(encoding='utf-8', errors='replace')
                    changed_files.append(_rel(locator_path, root))
                    created_symbols.append({'file': _rel(locator_path, root), 'kind': 'locator', 'name': locator, 'scenario_id': scenario.get('id'), 'evidence': 'provisional semantic locator; verify with Playwright MCP/codegen/live DOM before production use'})
            if re.search(rf'\b{re.escape(method)}\s*\(', text):
                reused_symbols.append({'file': _rel(page_path, root), 'kind': 'method', 'name': method, 'scenario_id': scenario.get('id')})
            else:
                locator_expression = str(binding.get('locator_expression') or 'this.{locator}').format(locator=locator)
                member = _method_member(action, method, locator_expression)
                if _append_member(page_path, member, method):
                    text += '\n' + member
                    changed_files.append(_rel(page_path, root))
                    created_symbols.append({'file': _rel(page_path, root), 'kind': 'method', 'name': method, 'scenario_id': scenario.get('id')})
            calls.append(_step_call(step, method))
        spec_path = _unique_spec_path(test_dir, feature, scenario, index)
        spec_path.parent.mkdir(parents=True, exist_ok=True)
        import_path = _relative_import(spec_path, page_path)
        title = str(scenario.get('title') or scenario.get('id') or f'{feature} {index}')
        source_id = str(scenario.get('id') or '')
        lines = [
            "import { test } from '@playwright/test';",
            f"import {{ {class_name} }} from '{import_path}';",
            '',
            f"test.describe({json.dumps(feature + ' - generated from approved testcase source')}, () => {{",
            f"  test({json.dumps(source_id + ' - ' + title)}, async ({{ page }}) => {{",
            f"    const screen = new {class_name}(page);",
        ]
        lines.extend('    ' + call for call in calls)
        lines.extend(['  });', '});', ''])
        spec_path.write_text('\n'.join(lines), encoding='utf-8')
        created_paths.add(_rel(spec_path, root))
        changed_files.append(_rel(spec_path, root))
        scenario_outputs.append({
            'scenario_id': scenario.get('id'),
            'scenario_title': title,
            'spec_file': _rel(spec_path, root),
            'page_file': _rel(page_path, root),
            'locator_file': _rel(locator_path, root),
            'bdd_source': bool(scenario.get('gherkin_source')),
        })

    changed_files = sorted(dict.fromkeys(changed_files))
    specs = [x['spec_file'] for x in scenario_outputs]
    validation = _validate_specs(root, specs, validate_generated, changed_files)
    attempted_changed_files = list(changed_files)
    attempted_specs = list(specs)
    rollback = {'performed': False, 'restored_files': [], 'deleted_created_files': [], 'reason': ''}
    if validation.get('ok') is False:
        rollback = _rollback_generation(root, backup_root, created_paths)
        changed_files = []
        specs = []
    plan = {
        'ok': validation.get('ok') is not False,
        'generated_at': datetime.now().isoformat(timespec='seconds'),
        'feature': feature,
        'framework_path': str(root),
        'testcase_file': str(testcase_path.relative_to(REPO_ROOT)),
        'scenario_count': len(scenarios),
        'generated_spec_count': len(specs),
        'attempted_spec_count': len(attempted_specs),
        'scenario_outputs': scenario_outputs,
        'changed_files': changed_files,
        'attempted_changed_files': attempted_changed_files,
        'rollback': rollback,
        'created_symbols': created_symbols,
        'reused_symbols': reused_symbols,
        'placement_preview': preview,
        'deep_framework_understanding_used': bool(deep.get('ok')),
        'provider_requested': provider,
        'model_requested': model,
        'bdd_output_mode': bdd_output_mode,
        'mcp_readiness': mcp_readiness,
        'mcp_codegen_policy': {
            'reuse_first': True,
            'live_locator_evidence': 'Existing locator/methods are reused first. New semantic locators are marked provisional and must be verified using prepared Playwright MCP/codegen/live DOM before production execution.',
            'codegen_command': f"npx --no-install playwright codegen {base_url}" if base_url else 'Set Application/base URL, then run npx --no-install playwright codegen <URL> from the selected framework.',
            'no_hidden_browser_interaction': True,
        },
        'backup_root': str(backup_root),
        'validation': validation,
        'policy': 'One normalized testcase/scenario produces one Playwright spec. Existing page files and methods are reused first; new support files are created only when no suitable file exists and the user permits it.',
    }
    json_report, html_report = _generate_report(root, plan)
    history = root / '.aiqa-history' / 'new-test-generation.jsonl'
    history.parent.mkdir(parents=True, exist_ok=True)
    with history.open('a', encoding='utf-8') as handle:
        handle.write(json.dumps(plan, ensure_ascii=False) + '\n')
    event_message = (f'Generated {len(specs)} Playwright spec(s) for {len(scenarios)} testcase(s)' if plan['ok'] else f'Validation failed; rolled back {len(attempted_specs)} attempted Playwright spec(s) and restored the framework')
    log_event('module2_existing_framework', event_message, status='done' if plan['ok'] else 'warning', progress=100, feature=feature, details={'changed_files': changed_files, 'attempted_changed_files': attempted_changed_files, 'validation': validation, 'rollback': rollback})
    return {
        'ok': plan['ok'],
        'generated_spec_count': len(specs),
        'attempted_spec_count': len(attempted_specs),
        'scenario_count': len(scenarios),
        'generated_specs': specs,
        'attempted_generated_specs': attempted_specs,
        'changed_files': changed_files,
        'attempted_changed_files': attempted_changed_files,
        'rollback': rollback,
        'created_symbols': created_symbols,
        'reused_symbols': reused_symbols,
        'generation_report': str(json_report),
        'generation_report_html': str(html_report),
        'extension_plan': plan,
        'message': (f"Generated {len(specs)} Playwright spec file(s) from {len(scenarios)} testcase(s) inside the selected existing framework. Exact changed files and reuse decisions are listed in the generation report." if plan['ok'] else f"Playwright validation failed for {len(attempted_specs)} attempted spec file(s). AstraHeal automatically rolled back all generated source/spec changes; review the validation evidence in the generation report before trying again."),
    }
